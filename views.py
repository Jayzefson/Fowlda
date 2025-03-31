from django.shortcuts import render, redirect, get_object_or_404
from django.views import View
from django.contrib import messages
from .forms import *
import PyPDF2
import zipfile
from docx import Document
import io
from django.http import HttpResponse
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import pandas as pd
from odf.opendocument import OpenDocumentSpreadsheet
from odf.table import Table, TableRow, TableCell
from odf.text import P
from moviepy.editor import VideoFileClip, AudioFileClip, ImageClip
import os
import tempfile
from PIL import Image
from django.contrib.auth import authenticate, login, logout
from django.db import IntegrityError

class IndexView(View):
    def get(self, request):
        form = CadastroForm()
        return render(request, 'index.html')
    
    def post(self, request):
        return render(request, 'index.html')

class ConversaoView(View):
    def get(self, request):
        form = ArquivoForm()
        log = User.objects.get(id=request.session.get('Logado')) if 'Logado' in request.session else None
        return render(request, 'conversao.html', {'form': form, 'usu': log})

    def post(self, request):
        form = ArquivoForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES['arquivo']
            formato_selecionado = request.POST.get('formato')
            base_name = uploaded_file.name.rsplit('.', 1)[0]

            try:
                if uploaded_file.name.endswith('.pdf'):
                    pdf_io = io.BytesIO(uploaded_file.read())
                    pdf_reader = PyPDF2.PdfReader(pdf_io)

                    if formato_selecionado == 'txt':
                        txt_content = [page.extract_text() for page in pdf_reader.pages]
                        response = HttpResponse("\n".join(txt_content), content_type='text/plain')
                        response['Content-Disposition'] = f'attachment; filename="{base_name}.txt"'

                    elif formato_selecionado == 'docx':
                        doc = Document()
                        for page in pdf_reader.pages:
                            text = page.extract_text()
                            if text:
                                doc.add_paragraph(text)

                        docx_io = io.BytesIO()
                        doc.save(docx_io)
                        docx_io.seek(0)
                        response = HttpResponse(docx_io, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        response['Content-Disposition'] = f'attachment; filename="{base_name}.docx"'

                elif uploaded_file.name.endswith('.docx'):
                    doc = Document(io.BytesIO(uploaded_file.read()))

                    if formato_selecionado == 'txt':
                        txt_content = [paragraph.text for paragraph in doc.paragraphs]
                        response = HttpResponse("\n".join(txt_content), content_type='text/plain')
                        response['Content-Disposition'] = f'attachment; filename="{base_name}.txt"'

                    elif formato_selecionado == 'pdf':
                        pdf_io = io.BytesIO()
                        c = canvas.Canvas(pdf_io, pagesize=letter)
                        width, height = letter

                        for paragraph in doc.paragraphs:
                            if height < 72:
                                c.showPage()
                                height = letter[1] - 72
                            c.drawString(72, height - 12, paragraph.text)
                            height -= 12

                        c.save()
                        pdf_io.seek(0)
                        response = HttpResponse(pdf_io, content_type='application/pdf')
                        response['Content-Disposition'] = f'attachment; filename="{base_name}.pdf"'

                elif uploaded_file.name.endswith('.txt'):
                    content = uploaded_file.read().decode('utf-8').splitlines()

                    if formato_selecionado == 'docx':
                        doc = Document()
                        for line in content:
                            doc.add_paragraph(line)

                        docx_io = io.BytesIO()
                        doc.save(docx_io)
                        docx_io.seek(0)
                        response = HttpResponse(docx_io, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        response['Content-Disposition'] = f'attachment; filename="{base_name}.docx"'

                    elif formato_selecionado == 'pdf':
                        pdf_io = io.BytesIO()
                        c = canvas.Canvas(pdf_io, pagesize=letter)
                        width, height = letter

                        for line in content:
                            if height < 72:
                                c.showPage()
                                height = letter[1] - 72
                            c.drawString(72, height - 12, line)
                            height -= 12

                        c.save()
                        pdf_io.seek(0)
                        response = HttpResponse(pdf_io, content_type='application/pdf')
                        response['Content-Disposition'] = f'attachment; filename="{base_name}.pdf"'

                else:
                    messages.error(request, "Formato de arquivo não suportado.")
                    return redirect('conversao')

                # Registrar no banco de dados
                Arquivo.objects.create(
                    arquivo=uploaded_file,
                    chave="alguma-chave-única",  # Atualizar com lógica adequada
                    usuario=request.user if request.user.is_authenticated else None
                )

                return response

            except Exception as e:
                messages.error(request, f"Erro ao processar o arquivo: {str(e)}")
                return redirect('conversao')

        return render(request, 'conversao.html', {'form': form, 'usu': request.session.get('Logado')})

    
class ConversaoPLAView(View):
    def get(self, request):
        form = ArquivoForm()
        if 'Logado' in request.session:
            log = User.objects.get(id=request.session.get('Logado'))
        else:
            log = None
        return render(request, 'conversaoPLA.html', {'form': form, 'usu':log})

    def post(self, request):
        form = ArquivoForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES['arquivo']
            base_name = uploaded_file.name.rsplit('.', 1)[0]
            form.save()

            if uploaded_file.name.endswith('.ods'):
                # Converter ODS para XLSX
                df = pd.read_excel(uploaded_file, engine='odf')
                xlsx_io = io.BytesIO()
                df.to_excel(xlsx_io, index=False, engine='openpyxl')
                xlsx_io.seek(0)

                response = HttpResponse(xlsx_io, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                response['Content-Disposition'] = f'attachment; filename="{base_name}.xlsx"'
                return response

            elif uploaded_file.name.endswith('.xlsx'):
                # Converter XLSX para ODS
                df = pd.read_excel(uploaded_file, engine='openpyxl')
                ods_io = io.BytesIO()

                # Criar um novo documento ODS
                doc = OpenDocumentSpreadsheet()
                table = Table(name="Sheet1")
                doc.spreadsheet.addElement(table)

                for row in df.itertuples(index=False):
                    table_row = TableRow()
                    for value in row:
                        cell = TableCell()
                        # Adicionar texto como elemento ODF
                        text_element = P(text=str(value))
                        cell.addElement(text_element)
                        table_row.addElement(cell)
                    table.addElement(table_row)

                doc.save(ods_io)
                ods_io.seek(0)

                response = HttpResponse(ods_io, content_type='application/vnd.oasis.opendocument.spreadsheet')
                response['Content-Disposition'] = f'attachment; filename="{base_name}.ods"'
                return response

            else:
                messages.error(request, "Formato de arquivo não suportado.")
                return redirect('conversaoPLA')

        return render(request, 'conversaoPLA.html', {'form': form, })
    
class ConversaoVIDView(View):
    def get(self, request):
        form = ArquivoForm()
        if 'Logado' in request.session:
            log = User.objects.get(id=request.session.get('Logado'))
        else:
            log = None
        return render(request, 'conversaoVID.html', {'form': form, 'usu':log})

    def post(self, request):
        form = ArquivoForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES['arquivo']
            base_name = uploaded_file.name.rsplit('.', 1)[0]
            form.save()

            if uploaded_file.name.endswith('.mp4'):
                # Converter MP4 para MP3
                with tempfile.NamedTemporaryFile(suffix='.mp4', delete=False) as temp_video_file:
                    for chunk in uploaded_file.chunks():
                        temp_video_file.write(chunk)
                    temp_video_file_name = temp_video_file.name

                # Cria o arquivo de áudio temporário
                audio_io = tempfile.NamedTemporaryFile(suffix='.mp3', delete=False)
                
                # Usa um contexto `with` para garantir que o vídeo esteja fechado antes de usar
                with VideoFileClip(temp_video_file_name) as video_clip:
                    audio_clip = video_clip.audio
                    audio_clip.write_audiofile(audio_io.name)
                    audio_clip.close()  # Fecha o clip de áudio

                # Garantindo que o arquivo de áudio seja fechado antes de abrir
                audio_io.close()

                response = HttpResponse(open(audio_io.name, 'rb'), content_type='audio/mpeg')
                response['Content-Disposition'] = f'attachment; filename="{base_name}.mp3"'
                
                # Limpeza de arquivos temporários
                os.remove(temp_video_file_name)
                os.remove(audio_io.name)
                return response

            else:
                messages.error(request, "Formato de arquivo não suportado.")
                return redirect('conversaoVID')

        return render(request, 'conversaoVID.html', {'form': form})
    
class ConversaoIMGView(View):
    def get(self, request):
        form = ArquivoForm()
        if 'Logado' in request.session:
            log = User.objects.get(id=request.session.get('Logado'))
        else:
            log = None
        return render(request, 'conversaoIMG.html', {'form': form, 'usu':log})

    def post(self, request):
        form = ArquivoForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES['arquivo']
            base_name = uploaded_file.name.rsplit('.', 1)[0]
            form.save()

            if uploaded_file.name.endswith('.jpg'):
                # Converter JPG para PNG
                with Image.open(uploaded_file) as img:
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as png_io:
                        img.save(png_io.name, format='PNG')
                        png_io.seek(0)
                        response = HttpResponse(open(png_io.name, 'rb'), content_type='image/png')
                        response['Content-Disposition'] = f'attachment; filename="{base_name}.png"'
                
                # Limpeza de arquivos temporários
                os.remove(png_io.name)
                return response

            elif uploaded_file.name.endswith('.png'):
                # Converter PNG para JPG
                with Image.open(uploaded_file) as img:
                    with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as jpg_io:
                        img.convert('RGB').save(jpg_io.name, format='JPEG')
                        jpg_io.seek(0)
                        response = HttpResponse(open(jpg_io.name, 'rb'), content_type='image/jpeg')
                        response['Content-Disposition'] = f'attachment; filename="{base_name}.jpg"'

                # Limpeza de arquivos temporários
                os.remove(jpg_io.name)
                return response

            else:
                messages.error(request, "Formato de arquivo não suportado.")
                return redirect('conversaoIMG')

        return render(request, 'conversaoIMG.html', {'form': form})

class HistoricoView(View):
    def get(self, request):
        logado_usuario = request.user if request.user.is_authenticated else None
        historico = []

        if logado_usuario:
            historico = Arquivo.objects.filter(usuario=logado_usuario).order_by('-tempo')[:15]
            
            request.session['Logado'] = logado_usuario.id

        form = CadastroForm()
        form2 = FeedForm()

        return render(request, 'historico.html', {
            'form': form,
            'logado': logado_usuario,
            'his': historico,
            'form2': form2,
            'usu': logado_usuario,
        })

    def post(self, request):
        if 'cadastro_form' in request.POST:
            form = CadastroForm(request.POST)
            if form.is_valid():
                try:
                    user = form.save()
                    login(request, user)
                    messages.success(request, 'Usuário cadastrado e logado com sucesso!')
                    return redirect('historico')
                except IntegrityError:
                    form.add_error('email', 'Este e-mail já está cadastrado.')

            return render(request, 'historico.html', {
                'form': form,
                'logado': request.user.is_authenticated,
                'his': [],
                'form2': FeedForm(),
            })

        elif 'feed_form' in request.POST:
            form2 = FeedForm(request.POST)
            if form2.is_valid():
                feed = form2.save(commit=False)
                feed.usuario = request.user
                feed.save()
                messages.success(request, 'Feedback enviado com sucesso!')
                return redirect('historico')

            messages.error(request, 'Erro ao enviar o feedback.')
            return render(request, 'historico.html', {
                'form': CadastroForm(),
                'logado': request.user.is_authenticated,
                'his': [],
                'form2': form2,
            })

def LoginView(request):
    if request.method == 'POST':
        email = request.POST.get('email')
        password = request.POST.get('password')

        # Verifica se há um usuário com este e-mail
        try:
            user = User.objects.get(email=email)
        except User.DoesNotExist:
            messages.error(request, 'Email ou senha inválidos.')
            return render(request, 'historico.html', {'error_message': 'Email ou senha inválidos.'})

        # Autentica com o username do usuário associado
        user = authenticate(request, username=user.username, password=password)

        if user is not None:
            login(request, user)
            messages.success(request, 'Login realizado com sucesso!')

            # Verifica se o usuário é superusuário
            if user.is_superuser:
                return redirect('historico')  # Redireciona para o painel de administração
            else:
                return redirect('historico')  # Redireciona para a página principal
        else:
            messages.error(request, 'Email ou senha inválidos.')
            return render(request, 'historico.html', {'error_message': 'Email ou senha inválidos.'})

    return render(request, 'historico.html')

def LogoutView(request):
    if request.method == 'POST':
        logout(request)
        messages.success(request, 'Logout realizado com sucesso!')
        return redirect('historico')

class AjudaView(View):
    def get(self, request):
        formatos = Formato.objects.all().order_by('categoria', 'extensao')  
        return render(request, 'ajuda.html', {'formatos': formatos})

    def post(self, request):
        return render(request, 'ajuda.html')

    
class FeedView(View):
    def get(self, request):
        feed = Feedback.objects.all().order_by('-data_feedback')
        return render(request, 'feed.html', {'feed': feed, 'resposta_form': RespostaFeedbackForm()})

    def post(self, request):
        if request.user.is_staff:
            form = RespostaFeedbackForm(request.POST)
            if form.is_valid():
                feedback_id = request.POST.get('feedback_id')
                feedback = Feedback.objects.get(id=feedback_id)
                feedback.resposta = form.cleaned_data['resposta']
                feedback.save()
                messages.success(request, 'Resposta enviada com sucesso!')
        return redirect('feed')

class MFeedView(View):
    def get(self, request):
        feedbacks = Feedback.objects.filter(usuario=request.user).order_by('-data_feedback')
        return render(request, 'meu_feedback.html', {'feedbacks': feedbacks, 'avaliacao_form': AvaliacaoFeedbackForm()})

    def post(self, request):
        form = AvaliacaoFeedbackForm(request.POST)
        if form.is_valid():
            feedback_id = request.POST.get('feedback_id')
            feedback = Feedback.objects.get(id=feedback_id, usuario=request.user)
            feedback.avaliacao = form.cleaned_data['avaliacao']
            feedback.save()
            messages.success(request, 'Avaliação salva com sucesso!')
        return redirect('meu_feedback')
    
class GerenciarView(View):
    def get(self, request):
        if not request.user.is_staff:  # Verifica se o usuário é admin
            messages.error(request, "Você não tem permissão para acessar esta página.")
            return redirect('home')  # Redireciona para a página inicial ou outra página

        usuarios = User.objects.all().order_by('username')
        formatos = Formato.objects.all().order_by('categoria', 'extensao')

        return render(request, 'gerenciar.html', {
            'usuarios': usuarios,
            'formatos': formatos,
        })

    def post(self, request):
        if not request.user.is_staff:
            messages.error(request, "Você não tem permissão para realizar esta ação.")
            return redirect('home')

        if 'usuario_id' in request.POST:
            # Excluir usuários
            usuario_id = request.POST.get('usuario_id')
            usuario = get_object_or_404(User, id=usuario_id)

            if usuario == request.user:
                messages.error(request, "Você não pode excluir sua própria conta.")
            else:
                usuario.delete()
                messages.success(request, f"Usuário {usuario.username} foi excluído com sucesso.")
        elif 'categoria' in request.POST:
            # Adicionar formato
            categoria = request.POST.get('categoria')
            extensao = request.POST.get('extensao')
            if categoria and extensao:
                Formato.objects.create(categoria=categoria, extensao=extensao)
                messages.success(request, f"Formato '{extensao}' adicionado à categoria '{categoria}'.")

        elif 'formato_id' in request.POST:
            # Excluir formatos
            formato_id = request.POST.get('formato_id')
            formato = get_object_or_404(Formato, id=formato_id)
            formato.delete()
            messages.success(request, f"Formato '{formato.extensao}' foi excluído com sucesso.")

        return redirect('gerenciar')
    
class DeleteHistorico(View):
    def get(self, request, id):
        try:
            historico = Arquivo.objects.get(id = id)
            historico.delete()
        except Arquivo.DoesNotExist:
            return redirect('historico')
        return redirect('historico')
    
class FeedbackDeleteView(View):
    def post(self, request):
        feedback_id = request.POST.get("feedback_id")
        feedback = get_object_or_404(Feedback, id=feedback_id)

        # Excluir o feedback
        feedback.delete()
        messages.success(request, "Feedback excluído com sucesso!")
        return redirect('feed')